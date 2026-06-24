import json
import os
from docx import Document

file_path = r"C:\Users\Император Безмолвный\Desktop\fill_race.json"
docx_path = r"C:\Users\Император Безмолвный\Desktop\result.docx"

if not os.path.exists(file_path):
    print("Ошибка: Файл не найден!")
else:
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            data = json.load(f)
        print("Содержимое файла:")
        print((json.dumps(data, ensure_ascii=False, indent=2)))     
    except Exception as e:
        print(f"Ошибка при чтении файла: {e}")

# Создаём документ
doc = Document()
doc.add_heading("Out_race", level=1)

for key, value in data.items():
    doc.add_paragraph(f"{key}: {value}")

doc.save(docx_path)
print("Документ создан:", docx_path)
